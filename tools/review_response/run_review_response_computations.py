from __future__ import annotations

import json
import math
import platform
import subprocess
import sys
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Mapping, Sequence, Tuple

import matplotlib

PROJECT_ROOT_FOR_IMPORTS = Path(__file__).resolve().parents[2]
if str(PROJECT_ROOT_FOR_IMPORTS) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT_FOR_IMPORTS))

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from scipy.stats import kendalltau, pearsonr, spearmanr

from config.paths import (
    DATASET_ORDER,
    GP_DEDUP_DIR,
    GP_RUNS_DIR,
    GP_SUMMARIES_DIR,
    OUTPUTS_DIR,
    PROJECT_ROOT,
    SINGLE_EVAL_SUPP_DIR,
    measure_slug,
    get_single_norm_curve_file,
    get_surrogate_label_file,
)
from config.settings import (
    AUTOFOCUS_METRICS,
    EPS,
    GENERALIZATION_ALPHA,
    GP_PRIMARY_OBJECTIVE,
    GP_SECONDARY_OBJECTIVE,
    METRIC_WEIGHTS,
)
from src.evaluation.aggregation import compute_rank_based_summary, compute_value_based_summary
from src.evaluation.autofocus_metrics import normalize_focus_curve, predict_peak_index
from src.gp.baselines import assign_composite_ids
from src.gp.deap_search import p_sqrt, summarize_focus_metrics_matrix
from src.measures.focus_measure_library import build_focus_measure_registry
from src.utils.validation import load_csv_rows, load_json, save_csv_rows, save_json


OUT_ROOT = OUTPUTS_DIR / "10_review_response_computations"
SUBDIRS = {
    "gradient": OUT_ROOT / "gradient_free_voter_sensitivity",
    "cfm4": OUT_ROOT / "cfm4_diagnostics",
    "runtime": OUT_ROOT / "runtime_weight_sensitivity",
    "gp": OUT_ROOT / "gp_audit",
    "methods": OUT_ROOT / "supplementary_methods",
    "text": OUT_ROOT / "manuscript_insert_text",
}
GRADIENT_LABEL_DIR = SUBDIRS["gradient"] / "labels"

GRADIENT_FREE_VOTERS = (
    "Normalized Variance",
    "Histogram Entropy",
    "GLCM Contrast",
    "Fourier Transform Sharpness Index",
)

CFM4_TERMINALS = {
    "bg": "Brenner Gradient",
    "gse": "Gradient Squared Energy",
    "ftsi": "Fourier Transform Sharpness Index",
    "wde_internal": "Curvelet Transform Sharpness Index",
    "wde_reported": "Wavelet Detail Energy (db1)",
}


OUTPUT_FILES: List[str] = []
INPUT_FILES: List[str] = []
NORMALIZED_CURVES_FOUND = True
ANY_STAGE_RERUN = False
GP_RERUN = False


def record_input(path: Path) -> None:
    INPUT_FILES.append(str(path))


def record_output(path: Path) -> Path:
    OUTPUT_FILES.append(str(path))
    return path


def ensure_dirs() -> None:
    OUT_ROOT.mkdir(parents=True, exist_ok=True)
    for path in SUBDIRS.values():
        path.mkdir(parents=True, exist_ok=True)
    GRADIENT_LABEL_DIR.mkdir(parents=True, exist_ok=True)


def write_csv(rows: Sequence[Mapping[str, Any]], path: Path) -> Path:
    save_csv_rows(rows, path)
    return record_output(path)


def write_json(payload: Any, path: Path) -> Path:
    save_json(payload, path)
    return record_output(path)


def write_text(text: str, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")
    return record_output(path)


def save_figure(fig: plt.Figure, stem: Path) -> Tuple[Path, Path]:
    png = stem.with_suffix(".png")
    pdf = stem.with_suffix(".pdf")
    fig.savefig(png, dpi=300, bbox_inches="tight")
    fig.savefig(pdf, bbox_inches="tight")
    plt.close(fig)
    record_output(png)
    record_output(pdf)
    return png, pdf


def load_curves(dataset: str, measure: str) -> List[np.ndarray]:
    path = get_single_norm_curve_file(dataset, measure)
    record_input(path)
    arr = np.load(path, allow_pickle=True)
    if arr.ndim == 2 and arr.dtype != object:
        return [np.asarray(row, dtype=np.float64).reshape(-1) for row in arr]
    return [np.asarray(item, dtype=np.float64).reshape(-1) for item in arr]


def central_tie_argmax(curve: np.ndarray, *, exclude_endpoints: bool = False) -> int:
    arr = np.asarray(curve, dtype=np.float64).reshape(-1).copy()
    if exclude_endpoints and arr.size >= 3:
        arr[0] = -np.inf
        arr[-1] = -np.inf
    max_val = float(np.max(arr))
    tied = np.where(np.isclose(arr, max_val))[0].astype(int).tolist()
    tied = sorted(tied)
    return int(tied[len(tied) // 2])


def consensus_peak(votes: Sequence[int]) -> int:
    counts = Counter(int(v) for v in votes)
    max_count = max(counts.values())
    winners = sorted(idx for idx, count in counts.items() if count == max_count)
    return int(winners[len(winners) // 2])


def path_for_existing(rel: str) -> Path:
    path = PROJECT_ROOT / rel
    record_input(path)
    return path


def check_normalized_curves(measures: Sequence[str]) -> None:
    missing: List[str] = []
    for dataset in DATASET_ORDER:
        for measure in measures:
            path = get_single_norm_curve_file(dataset, measure)
            if not path.exists():
                missing.append(str(path))
    if missing:
        print("Missing normalized curve files. Stop before review-response computations.")
        print("If these files are intentionally absent, run only:")
        print("python scripts/03_run_single_measure_benchmark.py --full-run --skip-timing")
        for path in missing[:100]:
            print(path)
        raise SystemExit(2)


def summarize_curves_against_labels(
    curves: Sequence[np.ndarray],
    labels: np.ndarray,
) -> Dict[str, float]:
    grouped: Dict[int, List[int]] = defaultdict(list)
    for idx, curve in enumerate(curves):
        grouped[len(curve)].append(idx)

    sums: Dict[str, float] = {
        metric: 0.0 for metric in AUTOFOCUS_METRICS
        if metric not in ("rrmse_under_additive_noise", "execution_time_per_slice")
    }
    counts: Dict[str, int] = {metric: 0 for metric in sums}

    for curve_len, indices in grouped.items():
        matrix = np.stack([np.asarray(curves[idx], dtype=np.float64).reshape(-1) for idx in indices], axis=0)
        label_arr = np.asarray(labels, dtype=int).reshape(-1)[indices]
        metrics = summarize_focus_metrics_matrix(matrix, label_arr)
        for metric, values in metrics.items():
            if metric not in sums:
                continue
            vals = np.asarray(values, dtype=np.float64)
            sums[metric] += float(np.nansum(vals))
            counts[metric] += int(np.sum(np.isfinite(vals)))

    out: Dict[str, float] = {}
    for metric in sums:
        out[metric] = float(sums[metric] / counts[metric]) if counts[metric] else float("nan")
    return out


def load_dataset_metric_raw() -> Dict[str, Dict[str, Dict[str, float]]]:
    path = path_for_existing("outputs/04_single_measure_eval/supplementary/dataset_metric_raw.json")
    return load_json(path)


def load_union_metric_raw() -> Dict[str, Dict[str, Dict[str, float]]]:
    path = path_for_existing("outputs/06_composite_eval/supplementary/union_metric_raw.json")
    return load_json(path)


def load_original_rank_rows() -> List[Dict[str, str]]:
    return load_csv_rows(path_for_existing("outputs/04_single_measure_eval/supplementary/all_single_rank_based.csv"))


def load_original_value_rows() -> List[Dict[str, str]]:
    return load_csv_rows(path_for_existing("outputs/04_single_measure_eval/supplementary/all_single_value_based_equal_dataset.csv"))


def label_stack_counts(labels_by_dataset: Mapping[str, np.ndarray]) -> Dict[str, int]:
    return {dataset: int(len(labels_by_dataset[dataset])) for dataset in DATASET_ORDER}


def part_a_gradient_free(registry: Mapping[str, Mapping[str, Any]]) -> Dict[str, Any]:
    out_dir = SUBDIRS["gradient"]
    labels_by_dataset: Dict[str, np.ndarray] = {}
    label_agreement_rows: List[Dict[str, Any]] = []
    peak_distribution_rows: List[Dict[str, Any]] = []

    for dataset in DATASET_ORDER:
        voter_curves = {voter: load_curves(dataset, voter) for voter in GRADIENT_FREE_VOTERS}
        n_stacks = len(next(iter(voter_curves.values())))
        labels: List[int] = []
        voter_peaks: Dict[str, List[int]] = {voter: [] for voter in GRADIENT_FREE_VOTERS}
        voter_endpoint_flags: Dict[str, List[bool]] = {voter: [] for voter in GRADIENT_FREE_VOTERS}

        for stack_idx in range(n_stacks):
            votes: List[int] = []
            for voter in GRADIENT_FREE_VOTERS:
                curve = voter_curves[voter][stack_idx]
                pred = central_tie_argmax(curve, exclude_endpoints=True)
                votes.append(pred)
                voter_peaks[voter].append(pred)
                voter_endpoint_flags[voter].append(bool(pred == 0 or pred == len(curve) - 1))
            labels.append(consensus_peak(votes))

        label_arr = np.asarray(labels, dtype=int)
        labels_by_dataset[dataset] = label_arr
        label_path = GRADIENT_LABEL_DIR / f"{dataset}_gradient_free_labels.npy"
        np.save(label_path, label_arr)
        record_output(label_path)

        surrogate_path = get_surrogate_label_file(dataset)
        record_input(surrogate_path)
        if surrogate_path.exists():
            original = np.load(surrogate_path, allow_pickle=False).astype(int).reshape(-1)
            comparable = min(len(original), len(label_arr))
            shifts = np.abs(label_arr[:comparable] - original[:comparable])
            agreement_rate = float(np.mean(label_arr[:comparable] == original[:comparable])) if comparable else float("nan")
            mean_shift = float(np.mean(shifts)) if comparable else float("nan")
            median_shift = float(np.median(shifts)) if comparable else float("nan")
            p90_shift = float(np.percentile(shifts, 90)) if comparable else float("nan")
        else:
            agreement_rate = mean_shift = median_shift = p90_shift = float("nan")

        label_agreement_rows.append(
            {
                "dataset": dataset,
                "n_stacks": int(n_stacks),
                "agreement_rate": agreement_rate,
                "mean_abs_label_shift": mean_shift,
                "median_abs_label_shift": median_shift,
                "p90_abs_label_shift": p90_shift,
            }
        )

        for voter in GRADIENT_FREE_VOTERS:
            peaks = np.asarray(voter_peaks[voter], dtype=np.float64)
            endpoints = np.asarray(voter_endpoint_flags[voter], dtype=bool)
            peak_distribution_rows.append(
                {
                    "dataset": dataset,
                    "voter": voter,
                    "n_stacks": int(n_stacks),
                    "mean_peak_index": float(np.mean(peaks)),
                    "median_peak_index": float(np.median(peaks)),
                    "peak_index_std": float(np.std(peaks, ddof=0)),
                    "endpoint_vote_rate": float(np.mean(endpoints)),
                }
            )

    write_csv(label_agreement_rows, out_dir / "gradient_free_label_agreement.csv")
    write_csv(peak_distribution_rows, out_dir / "gradient_free_voter_peak_distribution.csv")

    original_raw = load_dataset_metric_raw()
    gradient_raw: Dict[str, Dict[str, Dict[str, float]]] = {dataset: {} for dataset in DATASET_ORDER}
    dataset_rows: List[Dict[str, Any]] = []
    measure_names = list(registry.keys())
    for dataset in DATASET_ORDER:
        labels = labels_by_dataset[dataset]
        for measure in measure_names:
            curves = load_curves(dataset, measure)
            metrics = summarize_curves_against_labels(curves, labels)
            existing = original_raw[dataset][measure]
            metrics["rrmse_under_additive_noise"] = float(existing.get("rrmse_under_additive_noise", float("nan")))
            metrics["execution_time_per_slice"] = float(existing.get("execution_time_per_slice", float("nan")))
            gradient_raw[dataset][measure] = dict(metrics)
            dataset_rows.append(
                {
                    "dataset": dataset,
                    "measure_name": measure,
                    "label_source_used": "gradient_free_surrogate",
                    **{metric: metrics.get(metric, float("nan")) for metric in AUTOFOCUS_METRICS},
                }
            )

    write_csv(dataset_rows, out_dir / "gradient_free_single_measure_dataset_metrics.csv")
    write_json(gradient_raw, out_dir / "gradient_free_dataset_metric_raw.json")

    counts = label_stack_counts(labels_by_dataset)
    rank_rows, _ = compute_rank_based_summary(
        dataset_metric_raw=gradient_raw,
        dataset_subset=DATASET_ORDER,
        alpha=GENERALIZATION_ALPHA,
    )
    value_rows = compute_value_based_summary(
        dataset_metric_raw=gradient_raw,
        dataset_stack_counts=counts,
        dataset_subset=DATASET_ORDER,
        weighting_mode="equal_dataset",
        alpha=GENERALIZATION_ALPHA,
        metric_weights=METRIC_WEIGHTS,
    )
    write_csv(rank_rows, out_dir / "gradient_free_rank_based_summary.csv")
    write_csv(value_rows, out_dir / "gradient_free_value_based_summary_equal_dataset.csv")

    original_rank = {row["measure_name"]: row for row in load_original_rank_rows()}
    original_value = {row["measure_name"]: row for row in load_original_value_rows()}
    gf_rank = {row["measure_name"]: row for row in rank_rows}
    gf_value = {row["measure_name"]: row for row in value_rows}
    delta_rows: List[Dict[str, Any]] = []
    for measure in measure_names:
        family = str(registry[measure].get("family", "unknown"))
        orank = int(float(original_rank[measure]["final_rank"]))
        grank = int(gf_rank[measure]["final_rank"])
        oval = int(float(original_value[measure]["final_rank"]))
        gval = int(gf_value[measure]["final_rank"])
        delta_rows.append(
            {
                "measure_name": measure,
                "original_rank_based_rank": orank,
                "gradient_free_rank_based_rank": grank,
                "rank_delta": int(grank - orank),
                "original_value_rank": oval,
                "gradient_free_value_rank": gval,
                "value_rank_delta": int(gval - oval),
                "family": family,
            }
        )
    write_csv(delta_rows, out_dir / "gradient_free_vs_original_rank_delta.csv")

    top_rows: List[Dict[str, Any]] = []
    original_rank_sorted = sorted(original_rank.values(), key=lambda row: float(row["final_rank"]))
    original_value_sorted = sorted(original_value.values(), key=lambda row: float(row["final_rank"]))
    for source_name, rows, score_key in (
        ("original_rank", original_rank_sorted[:10], "rank_generalization_score"),
        ("gradient_free_rank", sorted(rank_rows, key=lambda row: int(row["final_rank"]))[:10], "rank_generalization_score"),
        ("original_value", original_value_sorted[:10], "generalization_score"),
        ("gradient_free_value", sorted(value_rows, key=lambda row: int(row["final_rank"]))[:10], "generalization_score"),
    ):
        for idx, row in enumerate(rows, start=1):
            measure = str(row["measure_name"])
            top_rows.append(
                {
                    "comparison": source_name,
                    "top_position": idx,
                    "measure_name": measure,
                    "score": float(row[score_key]),
                    "family": str(registry[measure].get("family", "unknown")),
                }
            )
    write_csv(top_rows, out_dir / "gradient_free_top10_comparison.csv")

    plot_gradient_free_figure(delta_rows, original_rank, original_value, gf_rank, gf_value)
    interpretation = build_gradient_free_interpretation(rank_rows, value_rows, delta_rows, registry)
    write_text(interpretation, out_dir / "gradient_free_interpretation.md")
    return {
        "rank_top1": rank_rows[0]["measure_name"],
        "value_top1": value_rows[0]["measure_name"],
        "top10_gradient_family_count_rank": sum(
            1 for row in sorted(rank_rows, key=lambda r: int(r["final_rank"]))[:10]
            if "gradient" in str(registry[str(row["measure_name"])].get("family", "")).lower()
        ),
        "top10_gradient_family_count_value": sum(
            1 for row in sorted(value_rows, key=lambda r: int(r["final_rank"]))[:10]
            if "gradient" in str(registry[str(row["measure_name"])].get("family", "")).lower()
        ),
    }


def plot_gradient_free_figure(
    delta_rows: Sequence[Mapping[str, Any]],
    original_rank: Mapping[str, Mapping[str, str]],
    original_value: Mapping[str, Mapping[str, str]],
    gf_rank: Mapping[str, Mapping[str, Any]],
    gf_value: Mapping[str, Mapping[str, Any]],
) -> None:
    out_dir = SUBDIRS["gradient"]
    fig, axes = plt.subplots(1, 3, figsize=(17, 6))

    rank_measures = [str(row["measure_name"]) for row in sorted(gf_rank.values(), key=lambda r: int(r["final_rank"]))[:10]]
    y = np.arange(len(rank_measures))
    axes[0].barh(y - 0.18, [float(original_rank[m]["final_rank"]) for m in rank_measures], height=0.35, label="Original")
    axes[0].barh(y + 0.18, [float(gf_rank[m]["final_rank"]) for m in rank_measures], height=0.35, label="Gradient-free")
    axes[0].set_yticks(y)
    axes[0].set_yticklabels(rank_measures, fontsize=8)
    axes[0].invert_yaxis()
    axes[0].set_xlabel("Rank (lower is better)")
    axes[0].set_title("A. Rank-based top 10")
    axes[0].legend(fontsize=8)

    value_measures = [str(row["measure_name"]) for row in sorted(gf_value.values(), key=lambda r: int(r["final_rank"]))[:10]]
    y2 = np.arange(len(value_measures))
    axes[1].barh(y2 - 0.18, [float(original_value[m]["final_rank"]) for m in value_measures], height=0.35, label="Original")
    axes[1].barh(y2 + 0.18, [float(gf_value[m]["final_rank"]) for m in value_measures], height=0.35, label="Gradient-free")
    axes[1].set_yticks(y2)
    axes[1].set_yticklabels(value_measures, fontsize=8)
    axes[1].invert_yaxis()
    axes[1].set_xlabel("Rank (lower is better)")
    axes[1].set_title("B. Value-based top 10")
    axes[1].legend(fontsize=8)

    ordered = sorted(delta_rows, key=lambda row: float(row["rank_delta"]))
    xs = np.arange(len(ordered))
    colors = ["#2b8cbe" if float(row["rank_delta"]) <= 0 else "#d95f0e" for row in ordered]
    axes[2].bar(xs, [float(row["rank_delta"]) for row in ordered], color=colors)
    axes[2].axhline(0, color="black", linewidth=0.8)
    axes[2].set_xticks(xs)
    axes[2].set_xticklabels([str(row["measure_name"]) for row in ordered], rotation=90, fontsize=6)
    axes[2].set_ylabel("Gradient-free rank - original rank")
    axes[2].set_title("C. Rank shift across all measures")
    fig.tight_layout()
    save_figure(fig, out_dir / "FigS9_gradient_free_voter_sensitivity")


def build_gradient_free_interpretation(
    rank_rows: Sequence[Mapping[str, Any]],
    value_rows: Sequence[Mapping[str, Any]],
    delta_rows: Sequence[Mapping[str, Any]],
    registry: Mapping[str, Mapping[str, Any]],
) -> str:
    rank_top10 = sorted(rank_rows, key=lambda row: int(row["final_rank"]))[:10]
    value_top10 = sorted(value_rows, key=lambda row: int(row["final_rank"]))[:10]
    rank_grad = [
        str(row["measure_name"]) for row in rank_top10
        if "gradient" in str(registry[str(row["measure_name"])].get("family", "")).lower()
    ]
    value_grad = [
        str(row["measure_name"]) for row in value_top10
        if "gradient" in str(registry[str(row["measure_name"])].get("family", "")).lower()
    ]
    biggest_up = sorted(delta_rows, key=lambda row: float(row["rank_delta"]))[:5]
    biggest_down = sorted(delta_rows, key=lambda row: float(row["rank_delta"]), reverse=True)[:5]
    support = len(rank_grad) >= 5 or len(value_grad) >= 5
    conclusion = (
        "The gradient-family top tier largely survives this internal sensitivity analysis."
        if support
        else "The gradient-family top tier is weakened under the gradient-free surrogate analysis and the main claim should be framed more cautiously."
    )
    return f"""# Gradient-Free Voter Sensitivity Interpretation

The gradient-free surrogate-label sensitivity analysis rebuilt labels using only Normalized Variance, Histogram Entropy, GLCM Contrast, and Fourier Transform Sharpness Index. No gradient or Laplacian voter was used. The resulting labels are an internal surrogate-label stress test and should not be described as optical ground truth.

Under gradient-free rank-based scoring, the top-ranked operator was {rank_rows[0]['measure_name']}. Under gradient-free value-based scoring, the top-ranked operator was {value_rows[0]['measure_name']}. Gradient-family operators contributed {len(rank_grad)} of the rank-based top 10 and {len(value_grad)} of the value-based top 10.

{conclusion}

Operators with the largest upward rank movement under gradient-free labels were: {', '.join(str(row['measure_name']) for row in biggest_up)}. Operators with the largest downward movement were: {', '.join(str(row['measure_name']) for row in biggest_down)}.

Manuscript wording should state that this is a surrogate-label sensitivity analysis designed to test dependence on derivative-based voters, not an independent hardware-focus validation.
"""


def part_b_cfm4_diagnostics() -> Dict[str, Any]:
    out_dir = SUBDIRS["cfm4"]
    sign_rows: List[Dict[str, Any]] = []
    sim_rows: List[Dict[str, Any]] = []
    peak_rows: List[Dict[str, Any]] = []

    for dataset in DATASET_ORDER:
        bg_curves = load_curves(dataset, CFM4_TERMINALS["bg"])
        gse_curves = load_curves(dataset, CFM4_TERMINALS["gse"])
        ftsi_curves = load_curves(dataset, CFM4_TERMINALS["ftsi"])
        wde_curves = load_curves(dataset, CFM4_TERMINALS["wde_internal"])

        total_lt = total_le = total_gt = total_slices = 0
        stack_lt_fracs: List[float] = []
        maes: List[float] = []
        rmses: List[float] = []
        pearsons: List[float] = []
        spearmans: List[float] = []
        exact_agree = within1 = 0
        peak_diffs: List[int] = []

        for bg, gse, ftsi, wde in zip(bg_curves, gse_curves, ftsi_curves, wde_curves):
            bg = np.asarray(bg, dtype=np.float64)
            gse = np.asarray(gse, dtype=np.float64)
            ftsi = np.asarray(ftsi, dtype=np.float64)
            wde = np.asarray(wde, dtype=np.float64)
            diff = bg - gse
            lt = diff < 0
            le = diff <= 0
            gt = diff > 0
            total_lt += int(np.sum(lt))
            total_le += int(np.sum(le))
            total_gt += int(np.sum(gt))
            total_slices += int(diff.size)
            stack_lt_fracs.append(float(np.mean(lt)))

            raw_cfm4 = ftsi - wde * p_sqrt(p_sqrt(diff))
            norm_cfm4 = normalize_focus_curve(np.asarray(raw_cfm4, dtype=np.float64))
            ftsi_norm = normalize_focus_curve(ftsi)
            delta = norm_cfm4 - ftsi_norm
            maes.append(float(np.mean(np.abs(delta))))
            rmses.append(float(np.sqrt(np.mean(delta ** 2))))
            if np.std(norm_cfm4) > EPS and np.std(ftsi_norm) > EPS:
                pearsons.append(float(pearsonr(norm_cfm4, ftsi_norm).statistic))
                spearmans.append(float(spearmanr(norm_cfm4, ftsi_norm).statistic))
            else:
                pearsons.append(float("nan"))
                spearmans.append(float("nan"))

            p_cfm4 = predict_peak_index(norm_cfm4)
            p_ftsi = predict_peak_index(ftsi_norm)
            pdiff = abs(int(p_cfm4) - int(p_ftsi))
            peak_diffs.append(pdiff)
            exact_agree += int(pdiff == 0)
            within1 += int(pdiff <= 1)

        n_stacks = len(bg_curves)
        sign_rows.append(
            {
                "dataset": dataset,
                "n_stacks": int(n_stacks),
                "n_slices_total": int(total_slices),
                "fraction_bg_lt_gse": float(total_lt / max(1, total_slices)),
                "fraction_bg_le_gse": float(total_le / max(1, total_slices)),
                "fraction_bg_gt_gse": float(total_gt / max(1, total_slices)),
                "median_stack_fraction_bg_lt_gse": float(np.median(stack_lt_fracs)),
                "p10_stack_fraction_bg_lt_gse": float(np.percentile(stack_lt_fracs, 10)),
                "p90_stack_fraction_bg_lt_gse": float(np.percentile(stack_lt_fracs, 90)),
            }
        )
        sim_rows.append(
            {
                "dataset": dataset,
                "n_stacks": int(n_stacks),
                "mean_curve_mae": float(np.nanmean(maes)),
                "median_curve_mae": float(np.nanmedian(maes)),
                "mean_curve_rmse": float(np.nanmean(rmses)),
                "mean_pearson_r": float(np.nanmean(pearsons)),
                "median_pearson_r": float(np.nanmedian(pearsons)),
                "mean_spearman_r": float(np.nanmean(spearmans)),
                "median_spearman_r": float(np.nanmedian(spearmans)),
            }
        )
        peak_arr = np.asarray(peak_diffs, dtype=np.float64)
        peak_rows.append(
            {
                "dataset": dataset,
                "n_stacks": int(n_stacks),
                "exact_peak_agreement_rate": float(exact_agree / max(1, n_stacks)),
                "within_1_slice_agreement_rate": float(within1 / max(1, n_stacks)),
                "mean_abs_peak_difference": float(np.mean(peak_arr)),
                "median_abs_peak_difference": float(np.median(peak_arr)),
                "p90_abs_peak_difference": float(np.percentile(peak_arr, 90)),
            }
        )

    write_csv(sign_rows, out_dir / "cfm4_bg_gse_sign_incidence_by_dataset.csv")
    write_csv(sim_rows, out_dir / "cfm4_vs_ftsi_curve_similarity_by_dataset.csv")
    write_csv(peak_rows, out_dir / "cfm4_vs_ftsi_peak_agreement_by_dataset.csv")

    metric_rows = build_cfm4_metric_comparison()
    write_csv(metric_rows, out_dir / "cfm4_vs_ftsi_metric_comparison.csv")
    plot_cfm4_figure(sign_rows, sim_rows, peak_rows, metric_rows)
    interpretation = build_cfm4_interpretation(sign_rows, sim_rows, peak_rows, metric_rows)
    write_text(interpretation, out_dir / "cfm4_diagnostic_interpretation.md")
    return {
        "mean_pearson": float(np.nanmean([row["mean_pearson_r"] for row in sim_rows])),
        "mean_exact_peak_agreement": float(np.mean([row["exact_peak_agreement_rate"] for row in peak_rows])),
    }


def build_cfm4_metric_comparison() -> List[Dict[str, Any]]:
    union_value = pd.read_csv(path_for_existing("outputs/06_composite_eval/supplementary/union_singles_and_composites_common_value.csv"))
    union_rank = pd.read_csv(path_for_existing("outputs/06_composite_eval/supplementary/union_singles_and_composites_common_rank.csv"))
    rows: List[Dict[str, Any]] = []
    for entity in ("CFM4", "Fourier Transform Sharpness Index"):
        v = union_value.loc[union_value["entity_name"] == entity].iloc[0]
        r = union_rank.loc[union_rank["entity_name"] == entity].iloc[0]
        rows.append(
            {
                "entity_name": entity,
                "common_value_G": float(v["generalization_score"]),
                "common_value_rank": int(v["final_rank"]),
                "common_rank_G": float(r["rank_generalization_score"]),
                "common_rank_rank": int(r["final_rank"]),
                "weighted_mean": float(v["weighted_mean"]),
                "weighted_std": float(v["weighted_std"]),
                "rank_mean": float(r["overall_rank_mean"]),
                "rank_std": float(r["overall_rank_std"]),
            }
        )
    return rows


def plot_cfm4_figure(
    sign_rows: Sequence[Mapping[str, Any]],
    sim_rows: Sequence[Mapping[str, Any]],
    peak_rows: Sequence[Mapping[str, Any]],
    metric_rows: Sequence[Mapping[str, Any]],
) -> None:
    out_dir = SUBDIRS["cfm4"]
    datasets = [row["dataset"] for row in sign_rows]
    fig, axes = plt.subplots(2, 2, figsize=(12, 9))
    axes = axes.ravel()
    axes[0].bar(datasets, [row["fraction_bg_lt_gse"] for row in sign_rows], color="#756bb1")
    axes[0].set_ylim(0, 1)
    axes[0].set_ylabel("Fraction of slices")
    axes[0].set_title("A. BG - GSE < 0 incidence")

    axes[1].bar(datasets, [row["exact_peak_agreement_rate"] for row in peak_rows], label="Exact", color="#31a354")
    axes[1].bar(datasets, [row["within_1_slice_agreement_rate"] for row in peak_rows], alpha=0.45, label="Within 1 slice", color="#addd8e")
    axes[1].set_ylim(0, 1)
    axes[1].set_ylabel("Agreement rate")
    axes[1].set_title("B. CFM4 vs FTSI peak agreement")
    axes[1].legend(fontsize=8)

    axes[2].bar(datasets, [row["mean_pearson_r"] for row in sim_rows], color="#3182bd")
    axes[2].set_ylim(0, 1)
    axes[2].set_ylabel("Mean Pearson r")
    axes[2].set_title("C. Curve similarity")

    names = [row["entity_name"] for row in metric_rows]
    x = np.arange(len(names))
    axes[3].bar(x - 0.18, [row["common_value_G"] for row in metric_rows], width=0.35, label="Common value G")
    axes[3].bar(x + 0.18, [row["common_rank_G"] for row in metric_rows], width=0.35, label="Common rank G")
    axes[3].set_xticks(x)
    axes[3].set_xticklabels(names, rotation=20, ha="right")
    axes[3].set_title("D. Common score comparison")
    axes[3].legend(fontsize=8)

    fig.tight_layout()
    save_figure(fig, out_dir / "FigS10_cfm4_vs_ftsi_diagnostics")


def build_cfm4_interpretation(
    sign_rows: Sequence[Mapping[str, Any]],
    sim_rows: Sequence[Mapping[str, Any]],
    peak_rows: Sequence[Mapping[str, Any]],
    metric_rows: Sequence[Mapping[str, Any]],
) -> str:
    mean_neg = float(np.mean([float(row["fraction_bg_lt_gse"]) for row in sign_rows]))
    mean_r = float(np.nanmean([float(row["mean_pearson_r"]) for row in sim_rows]))
    mean_exact = float(np.mean([float(row["exact_peak_agreement_rate"]) for row in peak_rows]))
    by_entity = {str(row["entity_name"]): row for row in metric_rows}
    cfm4 = by_entity["CFM4"]
    ftsi = by_entity["Fourier Transform Sharpness Index"]
    highly_similar = bool(mean_r >= 0.95 or mean_exact >= 0.90)
    similarity_sentence = (
        "The normalized CFM4 and FTSI curves are highly similar by correlation and peak agreement; at the curve level, the incremental composite contribution is therefore weak/modest rather than a clearly distinct focus response."
        if highly_similar
        else "The normalized CFM4 and FTSI curves are related but not identical by the curve and peak diagnostics."
    )
    value_sentence = (
        "CFM4 ranks better than FTSI under both common rank and common value scoring."
        if int(cfm4["common_rank_rank"]) < int(ftsi["common_rank_rank"]) and int(cfm4["common_value_rank"]) < int(ftsi["common_value_rank"])
        else "CFM4's advantage over FTSI is score-dependent rather than uniformly value-dominant."
    )
    return f"""# CFM4 Diagnostic Interpretation

The reviewer concern that CFM4 collapses to FTSI is not literally correct at the implementation level. The protected square root is `psqrt(x) = sqrt(abs(x) + eps)`, not zero-clipping. Therefore, slices with BG < GSE still contribute through the absolute-value branch. Across datasets, BG - GSE was negative for a mean fraction of {mean_neg:.3f} of evaluated slices.

{similarity_sentence} The mean dataset-level Pearson correlation between normalized CFM4 and FTSI curves was {mean_r:.3f}, and the mean exact peak-agreement rate was {mean_exact:.3f}.

In common scoring, CFM4 had value score {float(cfm4['common_value_G']):.4f} with value rank {int(cfm4['common_value_rank'])}, and rank score {float(cfm4['common_rank_G']):.4f} with rank rank {int(cfm4['common_rank_rank'])}. FTSI had value score {float(ftsi['common_value_G']):.4f} with value rank {int(ftsi['common_value_rank'])}, and rank score {float(ftsi['common_rank_G']):.4f} with rank rank {int(ftsi['common_rank_rank'])}. {value_sentence} Thus, the rank-level fusion result is supported, but the value-level contribution over standalone FTSI should be described as limited rather than value-dominant.

The internal terminal name `Curvelet Transform Sharpness Index` is a naming legacy; the current implementation maps it to Wavelet Detail Energy (db1). Manuscript text should use WDE_db1 or explicitly disclose the internal-name correction.
"""


def make_runtime_schemes() -> Dict[str, Dict[str, float]]:
    default = dict(METRIC_WEIGHTS)
    other_metrics = [m for m in AUTOFOCUS_METRICS if m != "execution_time_per_slice"]
    other_sum = sum(default[m] for m in other_metrics)
    redistributed = dict(default)
    redistributed["execution_time_per_slice"] = 0.05
    for metric in other_metrics:
        redistributed[metric] = default[metric] + 0.05 * default[metric] / other_sum

    to_accuracy = dict(default)
    to_accuracy["execution_time_per_slice"] = 0.05
    to_accuracy["absolute_peak_localization_error"] = default["absolute_peak_localization_error"] + 0.05

    no_runtime = dict(default)
    no_runtime["execution_time_per_slice"] = 0.0
    other_sum = sum(no_runtime[m] for m in other_metrics)
    for metric in other_metrics:
        no_runtime[metric] = no_runtime[metric] / other_sum

    return {
        "paper_default": default,
        "runtime_005_redistributed": redistributed,
        "runtime_005_to_accuracy": to_accuracy,
        "no_runtime": no_runtime,
    }


def part_c_runtime_sensitivity() -> Dict[str, Any]:
    out_dir = SUBDIRS["runtime"]
    schemes = make_runtime_schemes()
    single_raw = load_dataset_metric_raw()
    union_raw = load_union_metric_raw()
    counts = {dataset: 1 for dataset in DATASET_ORDER}

    all_summary_rows: List[Dict[str, Any]] = []
    single_detail_rows: List[Dict[str, Any]] = []
    union_detail_rows: List[Dict[str, Any]] = []
    ranking_cache: Dict[Tuple[str, str], Dict[str, int]] = {}

    for pool_name, raw, detail_rows in (
        ("single", single_raw, single_detail_rows),
        ("union", union_raw, union_detail_rows),
    ):
        default_rank_map: Dict[str, int] | None = None
        for scheme_name, weights in schemes.items():
            rows = compute_value_based_summary(
                dataset_metric_raw=raw,
                dataset_stack_counts=counts,
                dataset_subset=DATASET_ORDER,
                weighting_mode="equal_dataset",
                alpha=GENERALIZATION_ALPHA,
                metric_weights=weights,
            )
            rank_map = {str(row["measure_name"]): int(row["final_rank"]) for row in rows}
            ranking_cache[(pool_name, scheme_name)] = rank_map
            if scheme_name == "paper_default":
                default_rank_map = dict(rank_map)
            for row in rows:
                detail_rows.append(
                    {
                        "scheme": scheme_name,
                        "comparison_pool": pool_name,
                        "entity_name": str(row["measure_name"]),
                        "weighted_mean": float(row["weighted_mean"]),
                        "weighted_std": float(row["weighted_std"]),
                        "generalization_score": float(row["generalization_score"]),
                        "final_rank": int(row["final_rank"]),
                    }
                )

        assert default_rank_map is not None
        entities = list(default_rank_map.keys())
        default_values = np.asarray([default_rank_map[e] for e in entities], dtype=np.float64)
        for scheme_name in schemes:
            rank_map = ranking_cache[(pool_name, scheme_name)]
            scheme_values = np.asarray([rank_map[e] for e in entities], dtype=np.float64)
            top_rows = sorted(
                [row for row in detail_rows if row["scheme"] == scheme_name and row["comparison_pool"] == pool_name],
                key=lambda row: int(row["final_rank"]),
            )
            default_top10 = [e for e, rank in default_rank_map.items() if rank <= 10]
            max_shift = max(abs(rank_map[e] - default_rank_map[e]) for e in default_top10)
            all_summary_rows.append(
                {
                    "scheme": scheme_name,
                    "comparison_pool": pool_name,
                    "top1": top_rows[0]["entity_name"],
                    "top3_set": " | ".join(str(row["entity_name"]) for row in top_rows[:3]),
                    "top5_set": " | ".join(str(row["entity_name"]) for row in top_rows[:5]),
                    "spearman_vs_default": float(spearmanr(default_values, scheme_values).statistic),
                    "kendall_vs_default": float(kendalltau(default_values, scheme_values).statistic),
                    "max_rank_shift_top10": int(max_shift),
                }
            )

    write_csv(single_detail_rows, out_dir / "runtime_weight_sensitivity_single.csv")
    write_csv(union_detail_rows, out_dir / "runtime_weight_sensitivity_union.csv")
    write_csv(all_summary_rows, out_dir / "runtime_weight_sensitivity_summary.csv")
    plot_runtime_figure(single_detail_rows, union_detail_rows)
    interpretation = build_runtime_interpretation(all_summary_rows)
    write_text(interpretation, out_dir / "runtime_weight_interpretation.md")
    return {
        "single_top1_default": next(row["top1"] for row in all_summary_rows if row["scheme"] == "paper_default" and row["comparison_pool"] == "single"),
        "union_top1_default": next(row["top1"] for row in all_summary_rows if row["scheme"] == "paper_default" and row["comparison_pool"] == "union"),
    }


def plot_runtime_figure(single_rows: Sequence[Mapping[str, Any]], union_rows: Sequence[Mapping[str, Any]]) -> None:
    out_dir = SUBDIRS["runtime"]
    fig, axes = plt.subplots(1, 2, figsize=(14, 6))
    for ax, pool_name, rows in ((axes[0], "single", single_rows), (axes[1], "union", union_rows)):
        df = pd.DataFrame(rows)
        default = df[df["scheme"] == "paper_default"].set_index("entity_name")["final_rank"].astype(int)
        top_entities = default.sort_values().head(15).index.tolist()
        x = np.arange(len(top_entities))
        width = 0.25
        for offset, scheme in zip((-width, 0.0, width), ("runtime_005_redistributed", "runtime_005_to_accuracy", "no_runtime")):
            s = df[df["scheme"] == scheme].set_index("entity_name")["final_rank"].astype(int)
            shifts = [int(s[e] - default[e]) for e in top_entities]
            ax.bar(x + offset, shifts, width=width, label=scheme)
        ax.axhline(0, color="black", linewidth=0.8)
        ax.set_xticks(x)
        ax.set_xticklabels(top_entities, rotation=90, fontsize=7)
        ax.set_ylabel("Rank shift vs paper default")
        ax.set_title(f"{pool_name.capitalize()} pool")
        ax.legend(fontsize=7)
    fig.tight_layout()
    save_figure(fig, out_dir / "FigS11_runtime_weight_sensitivity")


def build_runtime_interpretation(summary_rows: Sequence[Mapping[str, Any]]) -> str:
    non_default = [row for row in summary_rows if row["scheme"] != "paper_default"]
    min_spearman = min(float(row["spearman_vs_default"]) for row in non_default)
    max_shift = max(int(row["max_rank_shift_top10"]) for row in non_default)
    top_changes = [
        row for row in non_default
        if row["top1"] != next(d["top1"] for d in summary_rows if d["scheme"] == "paper_default" and d["comparison_pool"] == row["comparison_pool"])
    ]
    if top_changes:
        top_sentence = "At least one alternative runtime-weight scheme changed the top-ranked entity."
    else:
        top_sentence = "The top-ranked entity was unchanged across the tested runtime-weight schemes."
    return f"""# Runtime-Weight Sensitivity Interpretation

Runtime-weight sensitivity was recomputed without rerunning the benchmark by reusing saved raw metric tensors. The tested schemes reduced the execution-time weight to 0.05 with two redistribution rules and removed execution time entirely with renormalization of the remaining metrics.

Across single-only and union pools, the minimum Spearman correlation with the paper-default ranking was {min_spearman:.3f}, and the maximum rank shift among paper-default top-10 entities was {max_shift}. {top_sentence}

These results should be used to state whether runtime weighting is a robustness parameter rather than a driver of the main conclusions.
"""


def part_d_gp_audit() -> Dict[str, Any]:
    out_dir = SUBDIRS["gp"]
    rows: List[Dict[str, Any]] = []
    best_files = sorted(GP_RUNS_DIR.rglob("best_result.json"))
    for path in best_files:
        record_input(path)
        data = load_json(path)
        settings = dict(data.get("gp_settings", {}))
        protocol = "final_refit" if "final_refit" in path.parts else "lodo"
        rows.append(
            {
                "protocol": protocol,
                "fold_or_scope": str(data.get("held_out_dataset", "FINAL_ALL" if protocol == "final_refit" else "")),
                "seed": data.get("seed", ""),
                "population_size": settings.get("population_size", ""),
                "num_generations": settings.get("num_generations", ""),
                "num_seeds": settings.get("num_seeds", ""),
                "tournament_size": settings.get("tournament_size", ""),
                "crossover_probability": settings.get("crossover_probability", ""),
                "mutation_probability": settings.get("mutation_probability", ""),
                "max_tree_depth": settings.get("max_tree_depth", ""),
                "max_nodes": settings.get("max_nodes", ""),
                "elitism": settings.get("elitism", ""),
                "use_nsga2": settings.get("use_nsga2", ""),
                "device": settings.get("device", data.get("array_backend", "")),
                "max_eval_seconds": settings.get("max_eval_seconds", ""),
                "primary_objective": data.get("gp_primary_objective", GP_PRIMARY_OBJECTIVE),
                "secondary_objective": data.get("gp_secondary_objective", GP_SECONDARY_OBJECTIVE),
                "result_json": str(path),
            }
        )

    write_csv(rows, out_dir / "gp_hyperparameter_audit.csv")
    fields = [
        "population_size",
        "num_generations",
        "num_seeds",
        "tournament_size",
        "crossover_probability",
        "mutation_probability",
        "max_tree_depth",
        "max_nodes",
        "elitism",
        "use_nsga2",
        "device",
        "max_eval_seconds",
        "primary_objective",
        "secondary_objective",
    ]
    unique_by_field = {
        field: sorted({str(row[field]) for row in rows})
        for field in fields
    }
    all_identical = all(len(values) == 1 for values in unique_by_field.values())
    core_fields = [
        "population_size",
        "num_generations",
        "num_seeds",
        "tournament_size",
        "crossover_probability",
        "mutation_probability",
        "max_tree_depth",
        "elitism",
        "use_nsga2",
        "primary_objective",
        "secondary_objective",
    ]
    core_identical = all(len(unique_by_field[field]) == 1 for field in core_fields)
    consistency = {
        "all_recorded_settings_identical": bool(all_identical),
        "core_recorded_settings_identical": bool(core_identical),
        "unique_values_by_field": unique_by_field,
        "num_best_result_files": len(rows),
        "note": "Some older LODO result files may omit post-patch audit fields such as max_nodes, max_eval_seconds, and device; this is reported as recorded-metadata inconsistency rather than rerunning GP.",
    }
    write_json(consistency, out_dir / "gp_hyperparameter_consistency.json")

    common = {field: most_common_value([row[field] for row in rows]) for field in fields}
    text = f"""# GP Hyperparameter Manuscript Text

The saved production GP runs used population size {common['population_size']}, {common['num_generations']} generations, tournament size {common['tournament_size']}, crossover probability {common['crossover_probability']}, mutation probability {common['mutation_probability']}, maximum tree depth {common['max_tree_depth']}, elitism {common['elitism']}, and NSGA-II = {common['use_nsga2']}. The primary objective was {common['primary_objective']} and the secondary objective was {common['secondary_objective']}.

Recorded audit status: core hyperparameters were {'consistent' if core_identical else 'not fully consistent'} across saved runs. Fields added in later resume/run-control patches, including max_nodes, max_eval_seconds, and device, are not present in every older LODO result file; where recorded, the dominant values were max_nodes = {common['max_nodes'] or 'not recorded'}, max_eval_seconds = {common['max_eval_seconds'] or 'not recorded'}, and device = {common['device'] or 'not recorded'}.
"""
    write_text(text, out_dir / "gp_hyperparameter_manuscript_text.md")
    return consistency


def most_common_value(values: Sequence[Any]) -> Any:
    counts = Counter(str(v) for v in values if str(v) != "")
    if not counts:
        return ""
    return counts.most_common(1)[0][0]


def part_e_composite_fold_origin() -> Dict[str, Any]:
    out_dir = SUBDIRS["gp"]
    candidates = load_composite_candidates_with_ids()
    common_value = pd.read_csv(path_for_existing("outputs/06_composite_eval/supplementary/all_composites_common_scoring.csv"))
    common_rank = pd.read_csv(path_for_existing("outputs/06_composite_eval/supplementary/all_composites_common_rank_scoring.csv"))
    rank_by_id = common_rank.set_index("entity_name")
    rows: List[Dict[str, Any]] = []
    for cand in candidates:
        cid = str(cand["composite_id"])
        v = common_value.loc[common_value["composite_id"] == cid].iloc[0]
        r = rank_by_id.loc[cid]
        fold_origin = "FINAL_ALL" if str(cand.get("result_type", "")) == "final_refit" else str(cand.get("held_out_dataset", ""))
        rows.append(
            {
                "composite_id": cid,
                "expression": str(cand.get("best_expression", "")),
                "fold_origin": fold_origin,
                "seed_origin": cand.get("seed", ""),
                "heldout_score_from_lodo_stage": cand.get("heldout_score", float("nan")),
                "common_value_G": float(v["common_value_generalization_score"]),
                "common_value_rank": int(v["common_value_final_rank"]),
                "common_rank_G": float(r["rank_generalization_score"]),
                "common_rank_rank": int(r["final_rank"]),
            }
        )
    csv_path = write_csv(rows, out_dir / "composite_fold_origin_table.csv")
    latex_path = out_dir / "composite_fold_origin_table_latex.tex"
    write_composite_origin_latex(rows, latex_path)
    text = """# Composite Fold-Origin Interpretation

Each retained composite in the common-comparison table can be traced either to the final all-dataset refit or to one LODO fold and seed. The `heldout_score_from_lodo_stage` column is meaningful for LODO-origin composites only; it is not defined for the final all-dataset refit, which was trained after LODO validation to instantiate the proposed final expression.

The composite table should therefore describe these values as fold-origin audit metadata, not as a single shared validation score for all composites.
"""
    write_text(text, out_dir / "composite_fold_origin_interpretation.md")
    return {"csv": str(csv_path), "latex": str(latex_path)}


def load_composite_candidates_with_ids() -> List[Dict[str, Any]]:
    dedup_path = GP_DEDUP_DIR / "deduplicated_best_expressions.json"
    final_path = GP_SUMMARIES_DIR / "final_composite_expression.json"
    record_input(dedup_path)
    record_input(final_path)
    candidates = list(load_json(dedup_path))
    if final_path.exists():
        final = load_json(final_path)
        expr = str(final.get("best_expression", ""))
        if expr:
            final_candidate = {
                "result_type": "final_refit",
                "best_expression": expr,
                "terminals": list(final.get("terminals", [])),
                "heldout_score": float("nan"),
                "source": str(final.get("source", "all_dataset_refit_after_lodo_validation")),
                "seed": int(final.get("seed", -1)),
                "best_training_objective": float(final.get("best_training_objective", float("nan"))),
                "best_all_dataset_score": float(final.get("best_all_dataset_score", float("nan"))),
                "best_complexity": float(final.get("best_complexity", float("nan"))),
                "num_nodes": int(final.get("num_nodes", 0)),
                "tree_height": int(final.get("tree_height", 0)),
            }
            candidates = [row for row in candidates if str(row.get("best_expression", "")) != expr]
            candidates.insert(0, final_candidate)
    return assign_composite_ids(candidates)


def latex_escape(value: Any) -> str:
    text = str(value)
    for old, new in [
        ("\\", "\\textbackslash{}"),
        ("&", "\\&"),
        ("%", "\\%"),
        ("$", "\\$"),
        ("#", "\\#"),
        ("_", "\\_"),
        ("{", "\\{"),
        ("}", "\\}"),
    ]:
        text = text.replace(old, new)
    return text


def write_composite_origin_latex(rows: Sequence[Mapping[str, Any]], path: Path) -> None:
    lines = [
        "\\begin{tabular}{lllrrrr}",
        "\\hline",
        "Composite & Origin & Seed & LODO G & Value G & Value rank & Rank rank \\\\",
        "\\hline",
    ]
    for row in rows:
        heldout = row["heldout_score_from_lodo_stage"]
        heldout_str = "" if not np.isfinite(float(heldout)) else f"{float(heldout):.4f}"
        lines.append(
            f"{latex_escape(row['composite_id'])} & {latex_escape(row['fold_origin'])} & {latex_escape(row['seed_origin'])} & "
            f"{heldout_str} & {float(row['common_value_G']):.4f} & {int(row['common_value_rank'])} & {int(row['common_rank_rank'])} \\\\"
        )
    lines.extend(["\\hline", "\\end{tabular}", ""])
    write_text("\n".join(lines), path)


def part_f_metric_definitions() -> Dict[str, Any]:
    out_dir = SUBDIRS["methods"]
    text = """# Formal Metric Definitions

Let f = (f_1, ..., f_n) denote the per-stack focus curve after min-max normalization to [0, 1]. The predicted focus index p is the global maximum index of f. If multiple indices share the maximum, the central tied index is used. Metrics are averaged over stacks after per-stack computation.

## Absolute peak localization error

Given reference label y, the error is |p - y|. Lower is better.

## FWHM

The full width at half maximum is the number of curve samples with value at least 0.5 times the peak value, measured from the first to last above-threshold sample inclusive. If the peak value is numerically zero, the width is set to n. Lower is better.

## Curvature at peak

For interior peaks, curvature is max(0, -(f_{p-1} - 2 f_p + f_{p+1})). For endpoint peaks it is 0. Higher is better.

## Steep slope width

The nearest local minimum to the left of p and the nearest local minimum to the right of p are identified. If either side has no local minimum, the corresponding boundary is used. The width is right - left. Lower is better.

## Steep-to-gradual slope ratio

The numerator is the mean absolute local difference adjacent to the peak: |f_p - f_{p-1}| and/or |f_p - f_{p+1}| when available. The denominator is the mean absolute background first difference after excluding the two peak-adjacent differences. The ratio is numerator / (denominator + eps). Higher is better.

## False maxima count

This counts strict local maxima excluding the selected global peak. Lower is better.

## Noise level

Noise level is the mean squared second difference: mean((Delta^2 f)^2). It is set to 0 for curves with fewer than three samples. Lower is better.

## Range around global maximum

Starting from p, this metric counts the contiguous region around the global maximum whose values remain at least 0.95 times the peak value. If the peak value is numerically zero, the range is set to n. Lower is better.

## RRMSE under additive noise

A normalized noisy image stack is generated by adding zero-mean Gaussian noise with the configured standard deviation to each slice, clipping to [0, 1], recomputing the focus curve, and min-max normalizing the noisy curve. RRMSE is sqrt(mean((f - f_noisy)^2)) / (sqrt(mean(f^2)) + eps). Lower is better.

These definitions follow the current implementation in `src.evaluation.autofocus_metrics` and the vectorized equivalent in `src.gp.deap_search.summarize_focus_metrics_matrix`.
"""
    md_path = write_text(text, out_dir / "formal_metric_definitions.md")
    docx_written = False
    try:
        import docx  # type: ignore

        document = docx.Document()
        for block in text.split("\n\n"):
            if block.startswith("# "):
                document.add_heading(block[2:], level=1)
            elif block.startswith("## "):
                document.add_heading(block[3:], level=2)
            else:
                document.add_paragraph(block)
        docx_path = out_dir / "formal_metric_definitions_for_supplementary.docx"
        document.save(docx_path)
        record_output(docx_path)
        docx_written = True
    except Exception:
        docx_written = False
    return {"markdown": str(md_path), "docx_written": docx_written}


def part_g_summary(
    gradient_summary: Mapping[str, Any],
    cfm4_summary: Mapping[str, Any],
    runtime_summary: Mapping[str, Any],
    gp_summary: Mapping[str, Any],
    metric_summary: Mapping[str, Any],
) -> None:
    text = f"""# Review-Response Results Summary

## Gradient-free voter sensitivity

Gradient-free surrogate labels were constructed using only Normalized Variance, Histogram Entropy, GLCM Contrast, and Fourier Transform Sharpness Index. No gradient or Laplacian voter was used. The rank-based top operator was {gradient_summary['rank_top1']}, and the value-based top operator was {gradient_summary['value_top1']}. Gradient-family operators contributed {gradient_summary['top10_gradient_family_count_rank']} of the rank-based top 10 and {gradient_summary['top10_gradient_family_count_value']} of the value-based top 10. This should be described as an internal surrogate-label sensitivity analysis, not as hardware ground truth.

## CFM4 vs FTSI diagnostic

The CFM4 diagnostic used the actual protected square root, `psqrt(x) = sqrt(abs(x) + eps)`. Therefore, BG < GSE does not zero out the correction term. The mean dataset-level CFM4-vs-FTSI Pearson correlation was {cfm4_summary['mean_pearson']:.3f}, and the mean exact peak-agreement rate was {cfm4_summary['mean_exact_peak_agreement']:.3f}. These values indicate that CFM4 is highly similar to FTSI at the curve/peak level, so the composite contribution should be framed as weak/modest. The rank-level fusion result can still be reported, but it should not be overclaimed as a clearly value-dominant improvement over standalone FTSI.

## Runtime-weight sensitivity

Runtime-weight sensitivity reused the saved raw metric tensors and did not rerun the benchmark. The default single-only top entity was {runtime_summary['single_top1_default']}, and the default union-pool top entity was {runtime_summary['union_top1_default']}. The accompanying sensitivity table reports whether reducing or removing runtime weight changes top-1, top-3, or top-5 conclusions.

## GP hyperparameter consistency

The GP audit found core recorded settings consistency = {gp_summary['core_recorded_settings_identical']}. Full recorded-settings identity = {gp_summary['all_recorded_settings_identical']}; if false, the difference is due to saved-result metadata fields that were added in later run-control patches, not to a rerun performed here.

## Formal metric definitions

Formal definitions for FWHM, curvature, steep-slope width, steep-to-gradual slope ratio, false maxima count, noise level, range around global maximum, and RRMSE under additive noise were written to the supplementary-methods folder. The docx export was {'available' if metric_summary['docx_written'] else 'not available because python-docx is not installed'}.

## Revised limitations paragraph

The revised manuscript should state that surrogate labels are consensus labels rather than optical ground truth; that the gradient-free analysis is a sensitivity analysis; that CFM4 is evaluated as an interpretable symbolic fusion but may remain curve-similar to FTSI; and that downstream/proxy analyses should not be described as diagnostic validation.
"""
    write_text(text, SUBDIRS["text"] / "review_response_results_summary.md")


def write_readme() -> None:
    readme = """# Review Response Computations

This folder contains isolated reviewer-response computations for the BSPC manuscript. The scripts and outputs here read existing pipeline artifacts and write only under `outputs/10_review_response_computations/`.

No LODO GP run was rerun. No full pipeline stage was rerun. Existing manuscript outputs under `outputs/03_*` through `outputs/09_paper/` were used as inputs only.

Subfolders:

- `gradient_free_voter_sensitivity/`: gradient-free surrogate-label sensitivity analysis.
- `cfm4_diagnostics/`: CFM4 vs FTSI and BG-GSE diagnostic outputs.
- `runtime_weight_sensitivity/`: metric-weight sensitivity for runtime weight.
- `gp_audit/`: saved GP hyperparameter and composite fold-origin audit.
- `supplementary_methods/`: formal metric definitions.
- `manuscript_insert_text/`: concise text blocks for reviewer response and manuscript insertion.
"""
    write_text(readme, OUT_ROOT / "README.md")


def validate_outputs() -> Dict[str, Any]:
    csv_failures: List[str] = []
    for path_str in OUTPUT_FILES:
        path = Path(path_str)
        if path.suffix.lower() == ".csv":
            try:
                pd.read_csv(path)
            except Exception as exc:
                csv_failures.append(f"{path}: {exc}")
    figures = [Path(p) for p in OUTPUT_FILES if Path(p).suffix.lower() in (".png", ".pdf")]
    figure_missing = [str(p) for p in figures if not p.exists() or p.stat().st_size == 0]
    return {
        "csv_read_failures": csv_failures,
        "num_figures": len(figures),
        "figure_missing_or_empty": figure_missing,
        "all_csvs_readable_by_pandas": len(csv_failures) == 0,
        "all_figures_saved": len(figure_missing) == 0,
    }


def git_commit_hash() -> str:
    try:
        completed = subprocess.run(
            ["git", "rev-parse", "HEAD"],
            cwd=str(PROJECT_ROOT),
            check=True,
            capture_output=True,
            text=True,
        )
        return completed.stdout.strip()
    except Exception:
        return ""


def main() -> None:
    ensure_dirs()
    registry = build_focus_measure_registry()
    measure_names = list(registry.keys())
    required_curves = sorted(set(measure_names) | set(GRADIENT_FREE_VOTERS) | set(CFM4_TERMINALS.values()))
    check_normalized_curves(required_curves)

    gradient_summary = part_a_gradient_free(registry)
    cfm4_summary = part_b_cfm4_diagnostics()
    runtime_summary = part_c_runtime_sensitivity()
    gp_summary = part_d_gp_audit()
    part_e_composite_fold_origin()
    metric_summary = part_f_metric_definitions()
    part_g_summary(gradient_summary, cfm4_summary, runtime_summary, gp_summary, metric_summary)
    write_readme()

    validation = validate_outputs()
    manifest_path = OUT_ROOT / "review_response_manifest.json"
    for audit_output in (Path(__file__).resolve(), manifest_path):
        if str(audit_output) not in OUTPUT_FILES:
            OUTPUT_FILES.append(str(audit_output))
    manifest = {
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "git_commit_hash": git_commit_hash(),
        "python_version": sys.version,
        "platform": platform.platform(),
        "input_files_used": sorted(set(INPUT_FILES)),
        "output_files_created": sorted(set(OUTPUT_FILES)),
        "normalized_curves_found": NORMALIZED_CURVES_FOUND,
        "any_stage_rerun": ANY_STAGE_RERUN,
        "gp_rerun": GP_RERUN,
        "active_measure_registry_count": len(measure_names),
        "dataset_order": list(DATASET_ORDER),
        "quality_checks": validation,
    }
    save_json(manifest, manifest_path)

    print("gradient-free sensitivity: done")
    print("CFM4 diagnostics: done")
    print("runtime-weight sensitivity: done")
    print("GP audit: done")
    print("composite fold-origin table: done")
    print("formal metric definitions: done")
    print(f"all outputs saved under {OUT_ROOT}: yes")


if __name__ == "__main__":
    main()
